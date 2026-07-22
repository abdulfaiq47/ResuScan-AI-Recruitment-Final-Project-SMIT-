from __future__ import annotations

import io
import re
from dataclasses import dataclass

from pypdf import PdfReader

MIN_USABLE_CHARS = 220  # below this we assume extraction failed / scanned CV


@dataclass
class ExtractedDoc:
    file_name: str
    text: str
    pages: int
    chars: int
    method: str
    warning: str = ""

    @property
    def ok(self) -> bool:
        return self.chars >= MIN_USABLE_CHARS


# ─────────────────────────────────────────────────────────────────────────────
#  Cleaning
# ─────────────────────────────────────────────────────────────────────────────
# Every key MUST be a non-empty string: str.replace("", x) injects x between
# every single character, which silently shreds the entire document.
_REPLACEMENTS = {
    # typographic ligatures
    "\ufb01": "fi", "\ufb02": "fl", "\ufb00": "ff", "\ufb03": "ffi", "\ufb04": "ffl",
    # bullet glyphs, including the private-use codepoints Word emits for
    # Symbol/Wingdings list markers
    # (U+00B7 middle dot is deliberately absent \u2014 resumes use it as a
    #  separator, "Python \u00b7 SQL", not as a list marker)
    "\u2022": "\u2022 ", "\u25cf": "\u2022 ", "\u25aa": "\u2022 ",
    "\u25e6": "\u2022 ", "\u2023": "\u2022 ",
    "\uf0b7": "\u2022 ", "\uf0a7": "\u2022 ", "\uf0d8": "\u2022 ",
    "\uf0fc": "\u2022 ", "\uf076": "\u2022 ",
    # dashes, exotic spaces and zero-width junk
    "\u2013": "-", "\u2014": "-", "\u2012": "-",
    "\u00a0": " ", "\u2007": " ", "\u202f": " ", "\u2009": " ",
    "\u200b": "", "\ufeff": "",
}

# A single empty key here would corrupt every uploaded file, so fail loudly.
assert all(_REPLACEMENTS), "empty replacement key would corrupt every document"


def clean_text(raw: str) -> str:
    """Normalise whitespace/bullets and drop page furniture, keeping structure."""
    if not raw:
        return ""

    text = raw
    for bad, good in _REPLACEMENTS.items():
        text = text.replace(bad, good)

    # Rejoin words split across a line break by hyphenation: "engi-\nneer"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    # Kill "Page 3 of 7" style furniture.
    text = re.sub(r"(?im)^\s*page\s+\d+\s*(of\s+\d+)?\s*$", "", text)
    # Collapse runs of spaces/tabs but never newlines (section structure matters).
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Max one blank line.
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    # Strip trailing spaces per line.
    text = "\n".join(line.rstrip() for line in text.splitlines())
    # Drop non-printables that survive some PDF encoders.
    text = "".join(ch for ch in text if ch.isprintable() or ch in "\n\t")
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
#  Extractors
# ─────────────────────────────────────────────────────────────────────────────
def _extract_pypdf(data: bytes) -> tuple[str, int]:
    reader = PdfReader(io.BytesIO(data))
    pages = [(p.extract_text() or "") for p in reader.pages]
    return "\n\n".join(pages), len(reader.pages)


def _extract_pdfplumber(data: bytes) -> tuple[str, int]:
    import pdfplumber  # imported lazily — only needed on the fallback path

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [(page.extract_text() or "") for page in pdf.pages]
        return "\n\n".join(pages), len(pdf.pages)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
#  Public API
# ─────────────────────────────────────────────────────────────────────────────
def read_document(file_name: str, data: bytes) -> ExtractedDoc:
    """Extract + clean any supported upload. Never raises — reports via .warning."""
    name = file_name or "document"
    lower = name.lower()

    try:
        if lower.endswith(".txt"):
            text = clean_text(data.decode("utf-8", errors="ignore"))
            return ExtractedDoc(name, text, 1, len(text), "plain-text")

        if lower.endswith(".docx"):
            text = clean_text(_extract_docx(data))
            return ExtractedDoc(name, text, 1, len(text), "python-docx")

        # ── PDF path ────────────────────────────────────────────────────────
        method, pages = "pypdf", 0
        try:
            raw, pages = _extract_pypdf(data)
        except Exception:
            raw = ""

        text = clean_text(raw)
        if len(text) < MIN_USABLE_CHARS:
            try:
                raw2, pages2 = _extract_pdfplumber(data)
                text2 = clean_text(raw2)
                if len(text2) > len(text):
                    text, pages, method = text2, pages2, "pdfplumber"
            except Exception:
                pass

        doc = ExtractedDoc(name, text, pages, len(text), method)
        if not doc.ok:
            doc.warning = (
                "Almost no text could be extracted — this looks like a scanned "
                "or image-only PDF. Re-save it as a text PDF, or paste the text "
                "manually."
            )
        return doc

    except Exception as exc:  # noqa: BLE001
        return ExtractedDoc(name, "", 0, 0, "failed", f"Could not read file: {exc}")


def read_upload(uploaded_file) -> ExtractedDoc:
    """Convenience wrapper for a Streamlit UploadedFile."""
    return read_document(uploaded_file.name, uploaded_file.getvalue())


def guess_name_from_filename(file_name: str) -> str:
    """Fallback candidate name when the LLM can't find one in the document."""
    stem = re.sub(r"\.(pdf|txt|docx)$", "", file_name or "", flags=re.I)
    stem = re.sub(r"(?i)\b(resume|cv|curriculum|vitae|final|updated|copy|new)\b", " ", stem)
    stem = re.sub(r"[_\-.\d()]+", " ", stem)
    stem = re.sub(r"\s{2,}", " ", stem).strip()
    return stem.title() or "Unknown Candidate"
